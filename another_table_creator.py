import docx

# Create a new Word file
document = docx.Document()

# Add a table with 4 columns and 8 rows (including header row)
table = document.add_table(rows=8, cols=3)

# Set table header style
header_row = table.rows[0]
header_row.cells[0].text = "Epic"
header_row.cells[1].text = "User Story"
header_row.cells[2].text = "Points"
for cell in header_row.cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# Add data to the table
data = [
    ["Plant Care Initiatives", "Plant care tips", "3"],
    ["Plant Care Initiatives", "Low-maintenance options", "3"],
    ["Plant Care Initiatives", "Watering reminders", "2"],
    ["Plant Care Initiatives", "Return policy", "2"],
    ["Plant Care Initiatives", "Plant care tools", "2"],
    ["Bonsai Trees", "Bonsai Selection", "13"],
    ["Bonsai Trees", "Bonsai Styles", "8"],
    ["Bonsai Trees", "Bonsai Tools", "13"]
]

for i in range(1, 8):
    row = table.rows[i]
    row.cells[0].text = data[i-1][0]
    row.cells[1].text = data[i-1][1]
    row.cells[2].text = data[i-1][2]

# Save the document
document.save("table.docx")