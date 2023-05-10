import docx

# Load the input document
doc = docx.Document("NotBold.docx")

# Replace all occurrences of "Peta" with "Mahmoud Tawfeek"
for paragraph in doc.paragraphs:
    if "Peta" in paragraph.text:
        paragraph.text = paragraph.text.replace("Peta", "Mahmoud Tawfeek")

# Create a new document for output
output_doc = docx.Document()

# Loop through each paragraph in the modified document
for paragraph in doc.paragraphs:
    # Check if the paragraph contains a colon
    if ":" in paragraph.text:
        # Split the paragraph at the colon
        parts = paragraph.text.split(":", maxsplit=1)
        # Create a new paragraph and add the first part in bold
        new_paragraph = output_doc.add_paragraph()
        new_run = new_paragraph.add_run(parts[0].strip())
        new_run.bold = True
        # Add the second part after the colon
        new_paragraph.add_run(": " + parts[1].strip())
    else:
        # If the paragraph doesn't contain a colon, just add it to the output document
        output_doc.add_paragraph(paragraph.text)

# Save the modified document to a new file
output_doc.save("LastModified.docx")